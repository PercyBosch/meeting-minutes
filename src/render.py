from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

from .models import Minutes, ActionItem


def action_line(item: ActionItem) -> str:
    extra = " · ".join(x for x in [item.owner, item.due] if x)
    return f"{item.task} ({extra})" if extra else item.task


def to_docx(minutes: Minutes, path) -> Path:
    doc = Document()
    doc.add_heading(minutes.title, level=0)
    doc.add_paragraph(f"Date: {minutes.date}")
    if minutes.attendees:
        doc.add_paragraph("Attendees / speakers: " + ", ".join(minutes.attendees))

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(minutes.summary)

    if minutes.key_points:
        doc.add_heading("Key Points", level=1)
        for p in minutes.key_points:
            doc.add_paragraph(p, style="List Bullet")

    if minutes.topics:
        doc.add_heading("Topics Discussed", level=1)
        for t in minutes.topics:
            doc.add_paragraph(t, style="List Number")

    # Meeting-only sections — shown only when they actually have content.
    if minutes.decisions:
        doc.add_heading("Key Decisions", level=1)
        for d in minutes.decisions:
            doc.add_paragraph(d, style="List Bullet")

    if minutes.action_items:
        doc.add_heading("Action Items", level=1)
        for a in minutes.action_items:
            doc.add_paragraph(action_line(a), style="List Bullet")

    if minutes.next_steps:
        doc.add_heading("Next Steps", level=1)
        for n in minutes.next_steps:
            doc.add_paragraph(n, style="List Bullet")

    doc.save(str(path))
    return Path(path)


def to_pdf(minutes: Minutes, path) -> Path:
    styles = getSampleStyleSheet()
    story = [
        Paragraph(minutes.title, styles["Title"]),
        Paragraph(f"Date: {minutes.date}", styles["Normal"]),
    ]
    if minutes.attendees:
        story.append(
            Paragraph("Attendees / speakers: " + ", ".join(minutes.attendees), styles["Normal"])
        )

    def section(heading, items, numbered=False):
        if not items:
            return
        story.append(Spacer(1, 12))
        story.append(Paragraph(heading, styles["Heading2"]))
        for i, item in enumerate(items, 1):
            bullet = f"{i}. " if numbered else "• "
            story.append(Paragraph(bullet + item, styles["Normal"]))

    story.append(Spacer(1, 12))
    story.append(Paragraph("Summary", styles["Heading2"]))
    story.append(Paragraph(minutes.summary, styles["Normal"]))
    section("Key Points", minutes.key_points)
    section("Topics Discussed", minutes.topics, numbered=True)
    section("Key Decisions", minutes.decisions)
    section("Action Items", [action_line(a) for a in minutes.action_items])
    section("Next Steps", minutes.next_steps)

    SimpleDocTemplate(str(path), pagesize=A4).build(story)
    return Path(path)
