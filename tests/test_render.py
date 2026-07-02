from docx import Document

from src.models import Minutes, ActionItem
from src.render import to_docx, to_pdf, action_line


def _sample() -> Minutes:
    return Minutes(
        title="Weekly Sync",
        date="2026-07-02",
        attendees=["Percy"],
        summary="We discussed the launch.",
        key_points=["Launch is on track", "Costs were flagged"],
        decisions=["Ship Friday"],
        action_items=[ActionItem(task="Write notes", owner="Percy", due="Thu")],
        topics=["Launch"],
        next_steps=["Rollback plan"],
    )


def test_action_line_includes_owner_and_due():
    assert action_line(ActionItem(task="Do X", owner="Percy", due="Thu")) == "Do X (Percy · Thu)"


def test_action_line_omits_empty_fields():
    assert action_line(ActionItem(task="Do X")) == "Do X"


def test_to_docx_writes_openable_file_with_headings(tmp_path):
    out = tmp_path / "m.docx"
    to_docx(_sample(), out)
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "Weekly Sync" in texts
    assert "Summary" in texts
    assert "Key Points" in texts
    assert "Key Decisions" in texts
    assert "Action Items" in texts


def test_to_docx_hides_empty_meeting_sections(tmp_path):
    # Podcast-style output: summary + key points, but no decisions/actions/next steps.
    podcast = Minutes(
        title="Interview",
        date="2026-07-02",
        summary="A wide-ranging chat.",
        key_points=["Point one", "Point two"],
        topics=["Media"],
    )
    out = tmp_path / "p.docx"
    to_docx(podcast, out)
    texts = [p.text for p in Document(str(out)).paragraphs]
    assert "Key Points" in texts
    assert "Key Decisions" not in texts
    assert "Action Items" not in texts
    assert "Next Steps" not in texts


def test_to_pdf_writes_nonempty_file(tmp_path):
    out = tmp_path / "m.pdf"
    to_pdf(_sample(), out)
    assert out.exists()
    assert out.stat().st_size > 0
